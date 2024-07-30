from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
import csv


def set_font(para, font_name='Arial', size=Pt(9)):
    for run in para.runs:
        run.font.name = font_name
        run.font.size = size

def add_page_number(paragraph):
    # Create a new run
    run = paragraph.add_run("Page ")

    # Set the font size to 7 points
    run.font.size = Pt(7)
    run.font.name = 'Arial'


    # Add the page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    # Set the font size for the page number
    run = paragraph.add_run()
    run.font.size = Pt(7)
    run.font.name = 'Arial'

def set_table_borders(table):
    tbl = table._element
    tbl_pr = tbl.xpath('./w:tblPr')[0]
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') # Border width
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'CCCCCC') # Border color light grey
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)

def add_header(doc):
    # Add header with page number in the center
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = 1  # Center alignment
    add_page_number(header_paragraph)
    # Add two empty lines after the page number
    header_paragraph.add_run("\n\n")

    # Add the title and placeholders
    p1 = header.add_paragraph("Student Name ________________________________________________________________________________________")
    p1.paragraph_format.left_indent = Inches(-0.5)
    p1.paragraph_format.right_indent = Inches(-0.5)
    set_font(p1)

    # Add class, section, and roll number placeholders
    p2 = header.add_paragraph("Class _____________________________    Section _____________________  Roll Number _________________________")
    p2.paragraph_format.left_indent = Inches(-0.5)
    p2.paragraph_format.right_indent = Inches(-0.5)
    set_font(p2)

    p3 = header.add_paragraph("_______________________________________________________________________________________________________")
    p3.paragraph_format.left_indent = Inches(-1)
    p3.paragraph_format.right_indent = Inches(-1)
    set_font(p3, size=Pt(10))
    run = p3.runs[0]
    run.bold = True
    p3.add_run("\n")

def add_question_section(doc, question_num, rows):
    p = doc.add_paragraph(f"Q {question_num}")
    # p.paragraph_format.left_indent = Inches(-0.5)
    set_font(p, size=Pt(10))
    run = p.runs[0]
    run.bold = True
    
    table = doc.add_table(rows=rows, cols=1)
    set_table_borders(table)
    
    # Add two blank lines after the table
    doc.add_paragraph()
    doc.add_paragraph()

def add_single_word_question(doc, question_number):
    p = doc.add_paragraph(f"Q {question_number} _____")        
    set_font(p, size=Pt(10))
    run = p.runs[0]
    run.bold = True
    p.add_run("\n\n")


def get_first_word(s):
    return s.split()[0] if s else ''



def create_answer_sheet():
    # Create a new Document
    doc = Document()

    add_header(doc)

    # Add section header
    p = doc.add_paragraph("SECTION 1")
    p.paragraph_format.alignment = 1
    set_font(p, size=Pt(17))  # Set larger size for the header
    run = p.runs[0]
    run.bold = True
    p.add_run("\n\n")


    # Read questions from CSV
    with open('question_details.csv', newline='') as csvfile:
        reader = csv.reader(csvfile)
        next(reader)  # Skip the header row
        for row in reader:
            # question_num, question_type, rows = row   
            question_num = f"{row[2]}{row[3]}{row[4]}"  # Concatenate Question Number, Part, and Subpart
            question_type = row[5]  # Get Question Type
            rows = get_first_word(row[5])  # Get the first word of Question Type

            if question_type == "Single Word":                
                add_single_word_question(doc, question_num)
            else:
                add_question_section(doc, question_num, int(rows))
    

    # Save the document
    doc.save("Generated_Answer_Sheet.docx")

    return doc




# create_answer_sheet()




